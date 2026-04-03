"""Generate styled resume as PDF (via WeasyPrint) and optionally as DOCX."""

import re
from pathlib import Path

import fitz  # PyMuPDF — for page counting
import jinja2
import weasyprint
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import OUTPUT_DIR, TEMPLATES_DIR


def _bold_md(text):
    """Convert **text** markdown bold to <strong> tags."""
    return jinja2.Markup(re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', str(text)))


def _add_heading_styled(doc: Document, text: str, level: int = 1):
    """Add a styled section heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(11) if level == 2 else Pt(16)


def _add_bullet(doc: Document, text: str, bold_prefix: str = ""):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)


def _add_role_header(doc: Document, company: str, dates: str, title: str, location: str = ""):
    """Add company/role header with dates right-aligned."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(44, 62, 80)

    if dates:
        p.add_run("  |  ").font.size = Pt(10)
        run_dates = p.add_run(dates)
        run_dates.font.size = Pt(9.5)
        run_dates.font.color.rgb = RGBColor(100, 100, 100)

    # Title line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    run_title = p2.add_run(title)
    run_title.italic = True
    run_title.font.size = Pt(10)
    if location:
        p2.add_run(f"  |  {location}").font.size = Pt(9)


def generate_styled_docx(resume_content: dict, filename: str = "resume_styled.docx") -> Path:
    """Generate a professionally styled DOCX resume.

    Args:
        resume_content: Structured resume dict from the writer.
        filename: Output filename.

    Returns:
        Path to the generated DOCX.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    header = resume_content.get("header", {})

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(header.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(44, 62, 80)

    # Target title
    if header.get("target_title"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(0)
        run = title_p.add_run(header["target_title"])
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(52, 73, 94)

    # Contact line
    contact_parts = [
        header.get("location", ""),
        header.get("phone", ""),
        header.get("email", ""),
        header.get("linkedin", ""),
    ]
    contact_line = " | ".join(p for p in contact_parts if p)
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(6)
    run = contact_p.add_run(contact_line)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Divider
    doc.add_paragraph("_" * 80).runs[0].font.color.rgb = RGBColor(189, 195, 199)

    # Summary
    _add_heading_styled(doc, "Professional Summary", level=2)
    summary = resume_content.get("summary", "")
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

    # Signature Achievements
    achievements = resume_content.get("signature_achievements", [])
    if achievements:
        _add_heading_styled(doc, "Key Achievements", level=2)
        for item in achievements:
            _add_bullet(doc, item)

    # Core Competencies
    competencies = resume_content.get("competencies", [])
    if competencies:
        _add_heading_styled(doc, "Core Competencies", level=2)
        p = doc.add_paragraph(" | ".join(competencies))
        for run in p.runs:
            run.font.size = Pt(9.5)

    # Professional Experience
    experience = resume_content.get("experience", [])
    if experience:
        _add_heading_styled(doc, "Professional Experience", level=2)
        for role in experience:
            _add_role_header(
                doc,
                role.get("company", ""),
                role.get("dates", ""),
                role.get("title", ""),
                role.get("location", ""),
            )
            if role.get("progression"):
                prog_p = doc.add_paragraph(role["progression"])
                prog_p.paragraph_format.space_before = Pt(0)
                for run in prog_p.runs:
                    run.font.size = Pt(8.5)
                    run.italic = True
                    run.font.color.rgb = RGBColor(120, 120, 120)

            if role.get("scope_line"):
                scope_p = doc.add_paragraph(role["scope_line"])
                scope_p.paragraph_format.space_before = Pt(2)
                for run in scope_p.runs:
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(44, 62, 80)

            for bullet in role.get("bullets", []):
                _add_bullet(doc, bullet)

    # Early Career
    early = resume_content.get("early_career", [])
    if early:
        _add_heading_styled(doc, "Early Career", level=2)
        for item in early:
            _add_bullet(doc, item)

    # Volunteering
    vol = resume_content.get("volunteering", {})
    if vol and vol.get("org"):
        _add_heading_styled(doc, "Volunteering", level=2)
        _add_role_header(doc, vol["org"], vol.get("dates", ""), vol.get("role", ""), vol.get("location", ""))
        if vol.get("bullet"):
            _add_bullet(doc, vol["bullet"])

    # Education
    education = resume_content.get("education", [])
    if education:
        _add_heading_styled(doc, "Education", level=2)
        for edu in education:
            _add_role_header(doc, edu["institution"], edu.get("dates", ""), edu["degree"])
            if edu.get("gpa"):
                p = doc.add_paragraph(edu["gpa"])
                for run in p.runs:
                    run.font.size = Pt(9)
            for h in edu.get("highlights", []):
                _add_bullet(doc, h)

    # Achievements & Tools
    at = resume_content.get("achievements_and_tools", {})
    if at:
        _add_heading_styled(doc, "Achievements & Tools", level=2)
        if at.get("achievements"):
            p = doc.add_paragraph()
            p.add_run("Achievements: ").bold = True
            p.add_run(" | ".join(at["achievements"])).font.size = Pt(9.5)
        if at.get("tools"):
            p = doc.add_paragraph()
            p.add_run("Tools: ").bold = True
            p.add_run(" | ".join(at["tools"])).font.size = Pt(9.5)

    # Save
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    return output_path


def generate_styled_pdf(resume_content: dict, filename: str = "resume_styled.pdf") -> Path:
    """Render resume as a styled PDF using WeasyPrint with the styled template.

    Automatically scales fonts down if content exceeds 2 pages.
    """
    env = jinja2.Environment(
        loader=jinja2.FileSystemLoader(str(TEMPLATES_DIR)),
        autoescape=True,
    )
    env.filters["bold_md"] = _bold_md
    template = env.get_template("styled.html")
    html_content = template.render(**resume_content)

    output_path = OUTPUT_DIR / filename
    css_path = TEMPLATES_DIR / "styled.css"

    MAX_PAGES = 2
    SCALE_START = 1.0
    SCALE_STEP = 0.03
    SCALE_MIN = 0.78

    scale = SCALE_START
    while scale >= SCALE_MIN:
        scale_css = weasyprint.CSS(string=f":root {{ --scale: {scale}; }}")
        base_css = weasyprint.CSS(filename=str(css_path))
        html = weasyprint.HTML(string=html_content, base_url=str(TEMPLATES_DIR))
        html.write_pdf(str(output_path), stylesheets=[base_css, scale_css])

        doc = fitz.open(str(output_path))
        pages = len(doc)
        doc.close()

        if pages <= MAX_PAGES:
            break
        scale -= SCALE_STEP

    return output_path
